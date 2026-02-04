#!/usr/bin/env python3
"""
Script untuk membuat dokumen analisis ML Development Lifecycle dalam format .docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def create_ml_analysis_document():
    doc = Document()
    
    # ========== TITLE ==========
    title = doc.add_heading('Analisis Machine Learning Development Lifecycle (ML DLC)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Sistem Rekomendasi Perawatan Sepsis Menggunakan Ensemble SAC')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # ========== DAFTAR ISI ==========
    doc.add_heading('Daftar Isi', level=1)
    toc_items = [
        "1. Pendahuluan",
        "2. Business Understanding",
        "3. Data Understanding & Preparation",
        "4. Feature Engineering",
        "5. Model Architecture (Modeling)",
        "6. Training Pipeline",
        "7. Evaluation"
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ========== 1. PENDAHULUAN ==========
    doc.add_heading('1. Pendahuluan', level=1)
    
    doc.add_paragraph(
        "Dokumen ini menyajikan analisis komprehensif terhadap proses pengembangan model Machine Learning "
        "dalam codebase PRISM (Personalized Recommendation for Intensive care Sepsis Management). "
        "Analisis ini mengikuti kerangka ML Development Lifecycle (ML DLC) yang mencakup tahapan "
        "dari persiapan data hingga evaluasi model."
    )
    
    doc.add_heading('1.1 Tujuan Proyek', level=2)
    doc.add_paragraph(
        "Pengembangan sistem rekomendasi perawatan sepsis berbasis Reinforcement Learning "
        "yang mampu memberikan rekomendasi dosis vasopressor dan IV fluid yang optimal "
        "untuk pasien sepsis di ICU, dengan tujuan memaksimalkan tingkat survival."
    )
    
    doc.add_heading('1.2 Pendekatan', level=2)
    doc.add_paragraph(
        "Proyek ini menggunakan pendekatan Offline Reinforcement Learning dengan algoritma "
        "Soft Actor-Critic (SAC) yang diimplementasikan dalam konfigurasi Ensemble "
        "untuk meningkatkan robustness dan mengurangi variance dalam prediksi."
    )
    
    doc.add_page_break()
    
    # ========== 2. BUSINESS UNDERSTANDING ==========
    doc.add_heading('2. Business Understanding', level=1)
    
    doc.add_heading('2.1 Konteks Permasalahan', level=2)
    doc.add_paragraph(
        "Sepsis merupakan kondisi kritis yang memerlukan penanganan cepat dan tepat. "
        "Pemberian vasopressor dan IV fluid yang tidak optimal dapat menyebabkan "
        "komplikasi serius atau kematian. Sistem ML ini dikembangkan untuk membantu "
        "klinisi dalam pengambilan keputusan perawatan."
    )
    
    doc.add_heading('2.2 Justifikasi Pendekatan Reinforcement Learning', level=2)
    justifications = [
        ("Sequential Decision Making", 
         "Perawatan sepsis melibatkan serangkaian keputusan yang berurutan, "
         "di mana setiap tindakan mempengaruhi kondisi pasien di waktu berikutnya."),
        ("Delayed Rewards", 
         "Outcome (survival/death) baru diketahui setelah periode waktu tertentu, "
         "bukan langsung setelah pemberian obat."),
        ("Personalized Treatment", 
         "Setiap pasien memiliki respons berbeda terhadap perawatan, "
         "sehingga diperlukan pendekatan yang dapat mempelajari treatment optimal per individu."),
        ("Learning from Historical Data", 
         "Tersedia data retrospektif dari MIMIC-IV yang berisi ribuan trajectory pasien "
         "yang dapat digunakan untuk pembelajaran offline.")
    ]
    
    for title_text, desc in justifications:
        p = doc.add_paragraph()
        runner = p.add_run(f"• {title_text}: ")
        runner.bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    # ========== 3. DATA UNDERSTANDING & PREPARATION ==========
    doc.add_heading('3. Data Understanding & Preparation', level=1)
    
    doc.add_heading('3.1 Sumber Data', level=2)
    doc.add_paragraph(
        "Data berasal dari database MIMIC-IV yang berisi informasi medis pasien ICU. "
        "Dataset telah diproses untuk mengekstrak trajectory pasien sepsis."
    )
    
    doc.add_heading('3.2 Struktur Data', level=2)
    
    # Create table for features
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Kategori'
    hdr_cells[1].text = 'Fitur'
    hdr_cells[2].text = 'Dimensi'
    
    data_rows = [
        ('State Features', 'Vital signs, lab values, demographics', '37 dimensi'),
        ('Actions', 'IV Fluid (mL), Vasopressor (mcg/kg/min)', '2 dimensi'),
        ('Outcomes', 'Mortality 90-hari (Y90)', '1 dimensi'),
        ('Severity Score', 'SOFA Score', '1 dimensi'),
    ]
    
    for cat, feat, dim in data_rows:
        row_cells = table.add_row().cells
        row_cells[0].text = cat
        row_cells[1].text = feat
        row_cells[2].text = dim
    
    doc.add_paragraph()
    
    doc.add_heading('3.3 Preprocessing Pipeline', level=2)
    
    doc.add_paragraph("Preprocessing dilakukan dalam dua tahap utama:")
    
    doc.add_heading('3.3.1 Normalisasi Kolom Numerik (colnorm)', level=3)
    code_block = doc.add_paragraph()
    code_block.add_run(
        "colnorm = ['SOFA', 'age', 'Weight_kg', 'GCS', 'HR', 'SysBP', 'MeanBP', 'DiaBP', 'RR', 'Temp_C',\n"
        "           'Sodium', 'Chloride', 'Glucose', 'Calcium', 'Hb', 'WBC_count', 'Platelets_count',\n"
        "           'PTT', 'PT', 'Arterial_pH', 'paO2', 'paCO2', 'HCO3', 'Arterial_lactate', 'Shock_Index',\n"
        "           'PaO2_FiO2', 'cumulated_balance', 'CO2_mEqL', 'Ionised_Ca']"
    ).font.size = Pt(9)
    
    doc.add_paragraph(
        "Justifikasi: Z-score normalization diterapkan untuk memastikan semua fitur numerik "
        "memiliki mean=0 dan std=1, yang penting untuk stabilitas training neural network."
    )
    
    doc.add_heading('3.3.2 Transformasi Logaritmik (collog)', level=3)
    code_block2 = doc.add_paragraph()
    code_block2.add_run(
        "collog = ['SpO2', 'BUN', 'Creatinine', 'SGOT', 'Total_bili', 'INR', 'input_total', 'output_total']"
    ).font.size = Pt(9)
    
    doc.add_paragraph(
        "Justifikasi: Transformasi log(0.1 + x) diterapkan pada fitur dengan distribusi skewed "
        "untuk mengkompres range nilai dan mengurangi pengaruh outlier."
    )
    
    doc.add_heading('3.3.3 Action Preprocessing', level=3)
    doc.add_paragraph(
        "Dosis obat ditransformasi menggunakan kombinasi Log1p dan MinMax scaling ke range [-1, 1]:\n\n"
        "• IV Fluid: log1p(dose_mL), kemudian scaled ke [-1, 1] dengan IV_LOG_MAX = 8.294 (≈4000 mL)\n"
        "• Vasopressor: log1p(dose_mcg), kemudian scaled ke [-1, 1] dengan VASO_LOG_MAX = 1.099 (≈2.0 mcg)\n\n"
        "Justifikasi: Transformasi logaritmik menangani distribusi dosis yang sangat skewed "
        "(banyak dosis rendah, sedikit dosis tinggi), sementara MinMax scaling memastikan "
        "output berada dalam range tanh activation function."
    )
    
    doc.add_page_break()
    
    # ========== 4. FEATURE ENGINEERING ==========
    doc.add_heading('4. Feature Engineering', level=1)
    
    doc.add_heading('4.1 AutoEncoder untuk Dimensionality Reduction', level=2)
    
    doc.add_paragraph(
        "Salah satu komponen kunci dalam pipeline adalah penggunaan AutoEncoder "
        "untuk mereduksi dimensi state dari 37 menjadi 24 dimensi latent."
    )
    
    doc.add_heading('4.1.1 Arsitektur AutoEncoder', level=3)
    
    # Architecture table
    ae_table = doc.add_table(rows=1, cols=3)
    ae_table.style = 'Table Grid'
    hdr = ae_table.rows[0].cells
    hdr[0].text = 'Komponen'
    hdr[1].text = 'Layer'
    hdr[2].text = 'Keterangan'
    
    ae_layers = [
        ('Encoder', 'Linear(37→128)', 'Input layer'),
        ('', 'LayerNorm + Mish + Dropout(0.1)', 'Normalization & activation'),
        ('', 'Linear(128→64)', 'Hidden layer'),
        ('', 'LayerNorm + Mish', 'Normalization & activation'),
        ('', 'Linear(64→24) + LayerNorm', 'Latent space output'),
        ('Decoder', 'Linear(24→64)', 'Latent to hidden'),
        ('', 'LayerNorm + Mish', 'Normalization & activation'),
        ('', 'Linear(64→128)', 'Hidden expansion'),
        ('', 'LayerNorm + Mish', 'Normalization & activation'),
        ('', 'Linear(128→37)', 'Reconstruction output'),
    ]
    
    for comp, layer, desc in ae_layers:
        row = ae_table.add_row().cells
        row[0].text = comp
        row[1].text = layer
        row[2].text = desc
    
    doc.add_paragraph()
    
    doc.add_heading('4.1.2 Justifikasi Penggunaan AutoEncoder', level=3)
    ae_justifications = [
        "Dimensionality Reduction: Mereduksi 37 fitur menjadi 24 fitur latent yang lebih compact, "
        "mengurangi curse of dimensionality untuk RL agent.",
        
        "Denoising: Training dengan noise injection (noise_factor=0.1) membuat encoder robust "
        "terhadap noise dalam data klinis.",
        
        "Feature Learning: AutoEncoder mempelajari representasi yang meaningful dari data, "
        "mengkompres informasi penting dan mengabaikan noise.",
        
        "Normalisasi Output: LayerNorm pada output latent memastikan representasi ter-normalisasi, "
        "memudahkan training downstream model."
    ]
    
    for i, just in enumerate(ae_justifications, 1):
        doc.add_paragraph(f"{i}. {just}")
    
    doc.add_heading('4.1.3 Training AutoEncoder', level=3)
    training_details = [
        ("Optimizer", "AdamW dengan learning rate 1e-3"),
        ("Scheduler", "ReduceLROnPlateau (factor=0.5, patience=5)"),
        ("Loss Function", "MSE Loss (reconstruction error)"),
        ("Regularization", "Dropout 0.1, Early stopping (patience=10)"),
        ("Batch Size", "256"),
    ]
    
    for param, value in training_details:
        p = doc.add_paragraph()
        p.add_run(f"• {param}: ").bold = True
        p.add_run(value)
    
    doc.add_page_break()
    
    # ========== 5. MODEL ARCHITECTURE ==========
    doc.add_heading('5. Model Architecture (Modeling)', level=1)
    
    doc.add_heading('5.1 Overview Arsitektur', level=2)
    doc.add_paragraph(
        "Sistem menggunakan Ensemble Soft Actor-Critic (SAC) dengan 5 agent independen. "
        "Setiap agent terdiri dari Actor (policy network) dan dual Critic networks."
    )
    
    doc.add_heading('5.2 Komponen Model', level=2)
    
    # ---- 5.2.1 Behavior Cloning ----
    doc.add_heading('5.2.1 Behavior Cloning Model', level=3)
    doc.add_paragraph(
        "Model Behavior Cloning digunakan untuk mengestimasi physician policy (μ) "
        "yang diperlukan untuk importance sampling dalam evaluasi."
    )
    
    bc_arch = doc.add_paragraph()
    bc_arch.add_run("Arsitektur:\n").bold = True
    bc_arch.add_run(
        "• Input: State latent (24 dim)\n"
        "• Hidden: 2× Linear(256) + LayerNorm + ReLU + Dropout(0.1)\n"
        "• Output: Mean layer (2 dim) + Log-Std layer (2 dim, clamped [-4, 2])\n"
        "• Loss: Negative Log Likelihood dengan Tanh-corrected Jacobian"
    )
    
    doc.add_paragraph(
        "Justifikasi: BC model mempelajari distribusi action physician "
        "dengan output stochastic (mean, std), memungkinkan estimasi log probability "
        "yang diperlukan untuk importance weight calculation dalam OPE."
    )
    
    # ---- 5.2.2 Actor Network ----
    doc.add_heading('5.2.2 Actor Network (Policy)', level=3)
    
    actor_arch = doc.add_paragraph()
    actor_arch.add_run("Arsitektur:\n").bold = True
    actor_arch.add_run(
        "• Input: State latent (24 dim)\n"
        "• Hidden: 3× Linear(1024) + GELU + ReLU\n"
        "• Output: Mean layer + Log-Std layer (clamped [-20, 2])\n"
        "• Action: Sampled dari Normal(mean, std), kemudian Tanh squashing"
    )
    
    actor_just = [
        "GELU Activation: Lebih smooth dibanding ReLU, memungkinkan gradien mengalir lebih baik "
        "terutama untuk nilai negatif kecil.",
        
        "Wide Network (1024 hidden): Kapasitas besar untuk menangkap kompleksitas policy "
        "dalam domain medis yang rumit.",
        
        "Orthogonal Initialization: Mencegah vanishing/exploding gradients di awal training.",
        
        "Tanh Squashing: Membatasi output action ke [-1, 1], sesuai dengan normalized action space."
    ]
    
    doc.add_paragraph("Justifikasi Design Choices:")
    for just in actor_just:
        doc.add_paragraph(f"• {just}")
    
    # ---- 5.2.3 Critic Network ----
    doc.add_heading('5.2.3 Critic Network (Q-Function)', level=3)
    
    critic_arch = doc.add_paragraph()
    critic_arch.add_run("Arsitektur:\n").bold = True
    critic_arch.add_run(
        "• Input: Concatenate(State latent, Action) = 26 dim\n"
        "• Hidden: 4× Linear(1024) + Mish activation\n"
        "• Output: Single Q-value (scalar)"
    )
    
    critic_just = [
        "Dual Critic: Menggunakan 2 critic networks (Q1, Q2) dengan minimum diambil "
        "untuk mengatasi overestimation bias.",
        
        "Mish Activation: Lebih smooth dari ReLU, membantu gradient flow pada Q-function "
        "yang surface-nya sangat kompleks.",
        
        "Target Networks: Soft update (τ=0.005) untuk stabilitas TD learning."
    ]
    
    doc.add_paragraph("Justifikasi:")
    for just in critic_just:
        doc.add_paragraph(f"• {just}")
    
    # ---- 5.2.4 Ensemble SAC ----
    doc.add_heading('5.2.4 Ensemble SAC Agent', level=3)
    
    doc.add_paragraph(
        "Ensemble terdiri dari 5 SAC agent independen yang ditraining secara paralel."
    )
    
    ensemble_just = [
        "Uncertainty Quantification: Variance antar agent memberikan estimasi uncertainty "
        "pada prediksi action.",
        
        "Robustness: Ensemble mengurangi variance dan lebih robust terhadap distributional shift.",
        
        "Strategy Options: Tersedia beberapa strategi untuk mengombinasikan prediksi:\n"
        "  - Median: Default, robust terhadap outlier\n"
        "  - Mean: Rata-rata prediksi\n"
        "  - LCB (Lower Confidence Bound): Conservative, mean - 0.5×std"
    ]
    
    for just in ensemble_just:
        doc.add_paragraph(f"• {just}")
    
    doc.add_page_break()
    
    # ========== 5.3 SAC Algorithm ==========
    doc.add_heading('5.3 Soft Actor-Critic (SAC) Algorithm', level=2)
    
    doc.add_paragraph(
        "SAC adalah algoritma RL yang menggabungkan off-policy learning dengan entropy regularization."
    )
    
    doc.add_heading('5.3.1 Key Components', level=3)
    
    sac_components = [
        ("Entropy-Regularized Objective", 
         "Memaksimalkan expected return + entropy bonus untuk exploration. "
         "Formula: J = E[Σ γᵗ(r + α·H(π))]"),
        
        ("Automatic Temperature Tuning", 
         "Alpha (α) di-tune secara otomatis dengan target entropy = -action_dim. "
         "Memastikan balance antara exploitation dan exploration."),
        
        ("Twin Delayed Critics", 
         "Dua critic networks di-update setiap step, actor dan target networks "
         "di-update dengan soft update (τ=0.005)."),
         
        ("Clinician-Guided Learning", 
         "Menggunakan uncertainty-weighted clinician guidance: "
         "agent mengikuti physician lebih kuat saat SOFA score rendah dan uncertainty tinggi."),
    ]
    
    for title_text, desc in sac_components:
        p = doc.add_paragraph()
        p.add_run(f"• {title_text}: ").bold = True
        p.add_run(desc)
    
    doc.add_heading('5.3.2 Lagrangian Safety Constraint', level=3)
    
    doc.add_paragraph(
        "Implementasi menambahkan safety constraint menggunakan Lagrangian optimization:"
    )
    
    safety_details = [
        "Clinical Penalty: Menghukum pemberian dosis tinggi saat kondisi pasien sudah melewati threshold "
        "(MAP > 78 mmHg untuk vasopressor, Balance > 5000 mL untuk fluid).",
        
        "Lagrange Multiplier (λ): Di-optimize secara dinamis dengan target_safety = 0.15. "
        "Lambda di-clamp ke range [0.1, 10] untuk stabilitas.",
        
        "Actor Loss: L_actor = SAC_loss + λ·safety_penalty"
    ]
    
    for detail in safety_details:
        doc.add_paragraph(f"• {detail}")
    
    doc.add_heading('5.3.3 Justifikasi Pemilihan SAC', level=3)
    
    sac_why = [
        "Sample Efficiency: SAC adalah salah satu algoritma RL paling sample-efficient, "
        "penting untuk domain medis dengan data terbatas.",
        
        "Continuous Action Space: Mendukung action space kontinu (dosis obat), "
        "berbeda dengan DQN yang hanya discrete.",
        
        "Exploration: Entropy regularization mendorong exploration yang penting "
        "untuk menemukan treatment optimal yang mungkin berbeda dari physician.",
        
        "Offline RL Compatibility: Dengan modifikasi (conservative Q-learning spirit), "
        "SAC dapat digunakan untuk offline RL."
    ]
    
    for i, why in enumerate(sac_why, 1):
        doc.add_paragraph(f"{i}. {why}")
    
    doc.add_page_break()
    
    # ========== 6. TRAINING PIPELINE ==========
    doc.add_heading('6. Training Pipeline', level=1)
    
    doc.add_heading('6.1 Training Flow', level=2)
    
    training_steps = [
        ("1. AutoEncoder Pre-training", 
         "Train AutoEncoder terlebih dahulu pada data state selama 50 epochs "
         "untuk mendapatkan representasi latent yang baik."),
        
        ("2. Behavior Cloning Pre-training", 
         "Train BC model pada data (state, action) physician untuk mengestimasi "
         "physician policy yang digunakan dalam evaluasi."),
        
        ("3. SAC Agent Training", 
         "Train 5 SAC agents secara paralel pada trajectory data, "
         "menggunakan encoded state dari frozen AutoEncoder."),
        
        ("4. Hyperparameter Tuning", 
         "Learning rate scheduling dengan ReduceLROnPlateau, "
         "gradient clipping (max_norm=1.0) untuk stabilitas."),
    ]
    
    for step, desc in training_steps:
        p = doc.add_paragraph()
        p.add_run(f"{step}: ").bold = True
        p.add_run(desc)
    
    doc.add_heading('6.2 Hyperparameters', level=2)
    
    hp_table = doc.add_table(rows=1, cols=3)
    hp_table.style = 'Table Grid'
    hdr = hp_table.rows[0].cells
    hdr[0].text = 'Parameter'
    hdr[1].text = 'Value'
    hdr[2].text = 'Justifikasi'
    
    hyperparams = [
        ('Learning Rate (Actor/Critic)', '3e-4', 'Standard untuk continuous control'),
        ('Discount Factor (γ)', '0.99', 'Long-horizon untuk survival outcome'),
        ('Soft Update (τ)', '0.005', 'Smooth target network update'),
        ('Initial Alpha', '0.2', 'Standard entropy coefficient'),
        ('Batch Size', '128', 'Balance memory dan gradient noise'),
        ('Hidden Dim', '1024', 'Large capacity untuk domain kompleks'),
        ('Num Ensemble Agents', '5', 'Cukup untuk uncertainty estimation'),
    ]
    
    for param, val, just in hyperparams:
        row = hp_table.add_row().cells
        row[0].text = param
        row[1].text = val
        row[2].text = just
    
    doc.add_heading('6.3 Reward Design', level=2)
    
    doc.add_paragraph(
        "Reward function di-design untuk mencerminkan objective klinis:"
    )
    
    reward_formula = doc.add_paragraph()
    reward_formula.add_run("R(s,a,s') = ").bold = True
    reward_formula.add_run(
        "−β₀·SOFA(s') + β₁·(SOFA(s) − SOFA(s')) + R_terminal\n\n"
        "di mana:\n"
        "• β₀, β₁ = weighting factors untuk severity dan improvement\n"
        "• R_terminal = +reward_value (survive) atau −reward_value (death)"
    )
    
    doc.add_paragraph(
        "Justifikasi: Reward mencakup immediate feedback (perubahan SOFA score) "
        "dan delayed feedback (survival outcome), sesuai dengan sifat sekuensial "
        "dari treatment sepsis."
    )
    
    doc.add_page_break()
    
    # ========== 7. EVALUATION ==========
    doc.add_heading('7. Evaluation', level=1)
    
    doc.add_paragraph(
        "Evaluasi model menggunakan Off-Policy Evaluation (OPE) karena tidak memungkinkan "
        "untuk melakukan online evaluation pada pasien nyata."
    )
    
    doc.add_heading('7.1 Metrik Evaluasi', level=2)
    
    # ---- 7.1.1 Q-Value Comparison ----
    doc.add_heading('7.1.1 Q-Value Comparison', level=3)
    doc.add_paragraph(
        "Membandingkan Q-value dari action agent vs action physician:\n\n"
        "• Q_agent: E[Q(s, π_agent(s))]\n"
        "• Q_physician: E[Q(s, a_physician)]\n\n"
        "Jika Q_agent > Q_physician secara konsisten, agent diprediksi lebih baik."
    )
    
    # ---- 7.1.2 WIS ----
    doc.add_heading('7.1.2 Weighted Importance Sampling (WIS)', level=3)
    doc.add_paragraph(
        "WIS mengestimasi expected return dari target policy menggunakan data dari behavior policy:\n\n"
        "V^WIS = Σ(wᵢ·rᵢ) / Σ(wᵢ)\n"
        "wᵢ = π_agent(a|s) / μ_physician(a|s)\n\n"
        "Justifikasi: WIS adalah estimator unbiased dan self-normalizing, "
        "mengatasi variance dari ordinary IS."
    )
    
    # ---- 7.1.3 DR ----
    doc.add_heading('7.1.3 Doubly Robust (DR) Estimator', level=3)
    doc.add_paragraph(
        "DR menggabungkan Direct Method dan Importance Sampling:\n\n"
        "V^DR = V̂(s) + ρ·(r + γV̂(s') − Q̂(s,a))\n\n"
        "Keunggulan:\n"
        "• Konsisten jika salah satu dari Q-function atau importance weight benar\n"
        "• Variance lebih rendah dari pure IS\n"
        "• Robust terhadap model misspecification"
    )
    
    # ---- 7.1.4 Survival Estimation ----
    doc.add_heading('7.1.4 Survival Rate Estimation', level=3)
    
    survival_methods = [
        ("Episode-Level IS", 
         "Menghitung geometric mean importance weight per episode, "
         "kemudian estimasi survival rate dari weighted outcomes."),
        
        ("DR-Calibrated Survival", 
         "Melatih calibrated logistic regression sebagai Q-model, "
         "kemudian menggunakan DR formula untuk estimasi counterfactual survival. "
         "Ini adalah metrik PRIMARY yang direkomendasikan."),
    ]
    
    for method, desc in survival_methods:
        p = doc.add_paragraph()
        p.add_run(f"• {method}: ").bold = True
        p.add_run(desc)
    
    doc.add_heading('7.2 Confidence Interval', level=2)
    doc.add_paragraph(
        "Semua metrik evaluasi dilengkapi dengan 95% confidence interval "
        "menggunakan Bootstrap resampling (2000 samples). "
        "Ini penting untuk menilai statistical significance dari hasil."
    )
    
    doc.add_heading('7.3 Effective Sample Size (ESS)', level=2)
    doc.add_paragraph(
        "ESS menghitung 'effective' jumlah sampel setelah importance weighting:\n\n"
        "ESS = (Σwᵢ)² / Σ(wᵢ²)\n\n"
        "ESS rendah mengindikasikan distributional shift yang besar antara policy, "
        "yang dapat menyebabkan high variance pada estimasi."
    )
    
    doc.add_heading('7.4 Action Alignment Analysis', level=2)
    doc.add_paragraph(
        "Analisis tambahan untuk memahami seberapa berbeda agent dari physician:\n\n"
        "• Action Distance: ||a_agent − a_physician||\n"
        "• Out-of-Range Rate: Persentase action di luar 5th-95th percentile physician\n"
        "• Distance by SOFA Group: Apakah agent lebih aggressive pada kasus severe?"
    )
    
    doc.add_heading('7.5 Justifikasi Metode Evaluasi', level=2)
    
    eval_justifications = [
        ("Multiple OPE Methods", 
         "Menggunakan beberapa metode (WIS, DR, Direct) untuk cross-validation hasil. "
         "Tidak bergantung pada satu metrik saja karena masing-masing memiliki asumsi berbeda."),
        
        ("Bootstrap CI", 
         "Confidence interval penting untuk domain medis karena overclaim "
         "dapat berbahaya. CI membantu menilai reliability estimasi."),
        
        ("ESS Monitoring", 
         "ESS < 100 mengindikasikan hasil tidak reliable karena high variance. "
         "Ini early warning untuk potential issues."),
        
        ("Behavior Cloning for μ", 
         "Menggunakan learned behavior model alih-alih global Gaussian "
         "untuk estimasi physician policy yang lebih akurat."),
    ]
    
    for method, desc in eval_justifications:
        p = doc.add_paragraph()
        p.add_run(f"• {method}: ").bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    # ========== KESIMPULAN ==========
    doc.add_heading('Kesimpulan', level=1)
    
    doc.add_paragraph(
        "Dokumen ini telah menganalisis secara komprehensif proses pengembangan model ML "
        "dalam sistem PRISM, mencakup:"
    )
    
    conclusions = [
        "Business Understanding: Konteks sepsis dan justifikasi penggunaan RL",
        "Data Preparation: Preprocessing dengan normalisasi dan log transform",
        "Feature Engineering: AutoEncoder untuk dimensionality reduction (37→24)",
        "Model Architecture: Ensemble SAC dengan 5 agents, Actor-Critic networks",
        "Training: Lagrangian safety constraints, clinician-guided learning",
        "Evaluation: Multiple OPE methods (WIS, DR) dengan bootstrap CI"
    ]
    
    for i, conc in enumerate(conclusions, 1):
        doc.add_paragraph(f"{i}. {conc}")
    
    doc.add_paragraph(
        "\nSetiap tahap telah dilengkapi dengan justifikasi teknis yang menjelaskan "
        "alasan di balik pilihan desain yang dibuat. Model ini dirancang dengan "
        "mempertimbangkan safety, robustness, dan reliability yang diperlukan "
        "untuk aplikasi di domain klinis."
    )
    
    # ========== SAVE ==========
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        'Analisis_ML_Development_Lifecycle.docx'
    )
    doc.save(output_path)
    print(f"✓ Dokumen berhasil disimpan: {output_path}")
    return output_path

if __name__ == "__main__":
    create_ml_analysis_document()
